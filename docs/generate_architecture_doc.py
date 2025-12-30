"""
Generate Spread Eagle Architecture Documentation.

Creates a beautiful Word document with flowcharts and architecture diagrams.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_heading(doc, text, level=1):
    """Add a styled heading."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)  # Dark blue
    return heading


def create_flowchart_table(doc, title, steps):
    """Create a visual flowchart using a table."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    table = doc.add_table(rows=len(steps), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    colors = {
        'source': 'E8F5E9',      # Light green
        'python': 'E3F2FD',      # Light blue
        'storage': 'FFF3E0',     # Light orange
        'transform': 'F3E5F5',   # Light purple
        'output': 'FFEBEE',      # Light red
    }

    for i, (icon, text, category) in enumerate(steps):
        # Icon cell
        cell0 = table.cell(i, 0)
        cell0.text = icon
        cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell0, colors.get(category, 'FFFFFF'))

        # Text cell
        cell1 = table.cell(i, 1)
        cell1.text = text
        set_cell_shading(cell1, colors.get(category, 'FFFFFF'))

        # Arrow cell (except last row)
        cell2 = table.cell(i, 2)
        if i < len(steps) - 1:
            cell2.text = "⬇"
            cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(0.8)
        row.cells[1].width = Inches(4.5)
        row.cells[2].width = Inches(0.5)

    return table


def create_architecture_box(doc, title, items, color='E3F2FD'):
    """Create a styled box with items."""
    table = doc.add_table(rows=1 + len(items), cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    header = table.cell(0, 0)
    header.text = title
    header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraphs[0].runs[0].bold = True
    set_cell_shading(header, color)

    # Items
    for i, item in enumerate(items):
        cell = table.cell(i + 1, 0)
        cell.text = f"  • {item}"
        set_cell_shading(cell, 'FFFFFF')

    for row in table.rows:
        row.cells[0].width = Inches(5)

    doc.add_paragraph()


def main():
    doc = Document()

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    title = doc.add_heading('', 0)
    title_run = title.add_run('🦅 SPREAD EAGLE')
    title_run.font.size = Pt(48)
    title_run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('Sports Betting Analytics Platform')
    subtitle_run.font.size = Pt(24)
    subtitle_run.font.color.rgb = RGBColor(0x5c, 0x6b, 0x7a)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    tagline = doc.add_paragraph()
    tagline_run = tagline.add_run('Find the Tail. Ride the Edge.')
    tagline_run.font.size = Pt(16)
    tagline_run.font.italic = True
    tagline_run.font.color.rgb = RGBColor(0x6c, 0x75, 0x7d)
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # =========================================================================
    # EXECUTIVE SUMMARY
    # =========================================================================
    add_styled_heading(doc, '📋 Executive Summary', 1)

    doc.add_paragraph(
        "Spread Eagle is a modern data analytics platform designed to identify "
        "low-volatility betting opportunities in college basketball. By analyzing "
        "historical betting lines, game statistics, and team performance metrics, "
        "the platform identifies teams that consistently play close to betting lines—"
        "creating valuable 'tail' opportunities for teaser bets."
    )

    doc.add_paragraph()

    # Key metrics table
    metrics_table = doc.add_table(rows=2, cols=4)
    metrics_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    metrics = [
        ('31,056', 'Games'),
        ('49,153', 'Betting Lines'),
        ('55,062', 'Team Stats'),
        ('48,551', 'Player Stats'),
    ]

    for i, (value, label) in enumerate(metrics):
        cell = metrics_table.cell(0, i)
        cell.text = value
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(18)
        set_cell_shading(cell, '1a365d')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        cell2 = metrics_table.cell(1, i)
        cell2.text = label
        cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell2, 'E3F2FD')

    doc.add_page_break()

    # =========================================================================
    # ARCHITECTURE OVERVIEW
    # =========================================================================
    add_styled_heading(doc, '🏗️ System Architecture', 1)

    doc.add_paragraph(
        "Spread Eagle follows a modern ELT (Extract, Load, Transform) architecture "
        "with clear separation between data ingestion, storage, and transformation layers."
    )

    doc.add_paragraph()

    # Architecture diagram as ASCII art
    arch_text = """
┌─────────────────────────────────────────────────────────────────────────────┐
│                           SPREAD EAGLE ARCHITECTURE                          │
└─────────────────────────────────────────────────────────────────────────────┘

    ┌──────────────────┐
    │   📡 DATA SOURCE │
    │  ─────────────── │
    │  College BBall   │
    │  Data API        │
    │  (REST/JSON)     │
    └────────┬─────────┘
             │
             ▼
    ┌──────────────────┐         ┌──────────────────┐
    │   🐍 PYTHON      │────────▶│   📦 AWS S3      │
    │  ─────────────── │         │  ─────────────── │
    │  Ingestion       │         │  Raw Data Lake   │
    │  Scripts         │         │  (JSON/CSV/      │
    │  (requests +     │         │   Parquet)       │
    │   boto3)         │         └────────┬─────────┘
    └──────────────────┘                  │
                                          ▼
                                 ┌──────────────────┐
                                 │   🐘 POSTGRESQL  │
                                 │  ─────────────── │
                                 │  AWS RDS         │
                                 │  (Raw Schema)    │
                                 └────────┬─────────┘
                                          │
                                          ▼
                                 ┌──────────────────┐
                                 │   🔄 DBT         │
                                 │  ─────────────── │
                                 │  Transform &     │
                                 │  Model           │
                                 │  (SQL + Jinja)   │
                                 └────────┬─────────┘
                                          │
                                          ▼
                                 ┌──────────────────┐
                                 │   📊 ANALYTICS   │
                                 │  ─────────────── │
                                 │  Rolling Stats   │
                                 │  ATS Analysis    │
                                 │  Volatility      │
                                 │  Scores          │
                                 └──────────────────┘
"""

    arch_para = doc.add_paragraph()
    arch_run = arch_para.add_run(arch_text)
    arch_run.font.name = 'Consolas'
    arch_run.font.size = Pt(8)

    doc.add_page_break()

    # =========================================================================
    # DATA FLOW
    # =========================================================================
    add_styled_heading(doc, '🔄 Data Flow Pipeline', 1)

    doc.add_paragraph(
        "Data flows through the system in distinct stages, each handled by specialized tools:"
    )

    flow_steps = [
        ('📡', 'API EXTRACTION: College Basketball Data API serves game stats, betting lines, team/player data', 'source'),
        ('🐍', 'PYTHON INGESTION: Custom scripts with date-range pagination handle API limits (3000 records/request)', 'python'),
        ('📦', 'S3 RAW STORAGE: JSON files by season + consolidated CSV/Parquet for analytics', 'storage'),
        ('🐘', 'RDS STAGING: PostgreSQL staging tables receive incremental loads (CDC pattern)', 'storage'),
        ('🔀', 'UPSERT: INSERT ON CONFLICT merges new data into main tables with load_date tracking', 'python'),
        ('🔄', 'DBT TRANSFORM: SQL models create rolling stats, ATS calculations, volatility metrics', 'transform'),
        ('📊', 'ANALYTICS LAYER: Mart tables ready for visualization and betting analysis', 'output'),
    ]

    create_flowchart_table(doc, '📊 End-to-End Data Pipeline', flow_steps)

    doc.add_page_break()

    # =========================================================================
    # TOOL STACK
    # =========================================================================
    add_styled_heading(doc, '🛠️ Technology Stack', 1)

    # Python section
    add_styled_heading(doc, '🐍 Python (Ingestion Layer)', 2)
    create_architecture_box(doc, 'Python Libraries', [
        'requests — HTTP client for API calls',
        'boto3 — AWS SDK for S3 uploads',
        'pandas — Data manipulation and CSV/Parquet export',
        'psycopg2 — PostgreSQL database driver',
        'pydantic — Settings and configuration management',
    ], 'E3F2FD')

    doc.add_paragraph(
        "Python handles all data extraction from the College Basketball Data API. "
        "Custom pagination logic using date ranges overcomes the API's 3000-record limit. "
        "Scripts are modular and follow a consistent pattern for easy extension to other sports."
    )

    # AWS section
    add_styled_heading(doc, '☁️ AWS Infrastructure', 2)
    create_architecture_box(doc, 'AWS Services', [
        'S3 — Raw data lake (JSON, CSV, Parquet)',
        'RDS PostgreSQL — Relational data warehouse',
        'IAM — Access control and permissions',
        'Secrets Manager — Secure credential storage',
        'VPC — Network isolation and security',
    ], 'FFF3E0')

    doc.add_paragraph(
        "Infrastructure is managed with Terraform for reproducibility. "
        "S3 provides cheap, durable storage for raw data with lifecycle policies. "
        "RDS PostgreSQL (db.t4g.micro) offers a cost-effective relational store (~$12/month)."
    )

    # dbt section
    add_styled_heading(doc, '🔄 dbt (Transform Layer)', 2)
    create_architecture_box(doc, 'dbt Features', [
        'SQL + Jinja templating for reusable logic',
        'Incremental models for efficient processing',
        'Built-in testing and documentation',
        'Lineage tracking and dependency management',
        'Modular structure: staging → intermediate → marts',
    ], 'F3E5F5')

    doc.add_paragraph(
        "dbt transforms raw data into analytics-ready tables. "
        "Models calculate rolling statistics, ATS performance, and volatility scores. "
        "The DAG ensures transformations run in the correct order with full lineage tracking."
    )

    doc.add_page_break()

    # =========================================================================
    # DATA TABLES
    # =========================================================================
    add_styled_heading(doc, '📊 Data Schema', 1)

    doc.add_paragraph(
        "The database schema is organized into reference tables (dimensions) "
        "and transactional tables (facts):"
    )

    # Tables summary
    tables_data = [
        ('Table', 'Type', 'Primary Key', 'Records', 'Description'),
        ('conferences', 'Reference', 'id', '34', 'NCAA conferences'),
        ('venues', 'Reference', 'id', '979', 'Arenas and stadiums'),
        ('teams', 'Reference', 'id', '1,515', 'College basketball teams'),
        ('games', 'Fact', 'id', '31,056', 'Game results and metadata'),
        ('betting_lines', 'Fact', 'game_id, provider', '49,153', 'Spreads, totals, moneylines'),
        ('game_team_stats', 'Fact', 'game_id, team_id', '55,062', 'Team box scores per game'),
        ('game_player_stats', 'Fact', 'game_id, athlete_id', '222,594', 'Player box scores'),
        ('team_season_stats', 'Fact', 'team_id, season', '3,515', 'Season aggregates by team'),
        ('player_season_stats', 'Fact', 'athlete_id, team_id, season', '48,551', 'Season aggregates by player'),
    ]

    table = doc.add_table(rows=len(tables_data), cols=5)
    table.style = 'Table Grid'

    for i, row_data in enumerate(tables_data):
        for j, cell_data in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = cell_data
            if i == 0:  # Header row
                set_cell_shading(cell, '1a365d')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].bold = True
            elif j == 1 and cell_data == 'Reference':
                set_cell_shading(cell, 'E8F5E9')
            elif j == 1 and cell_data == 'Fact':
                set_cell_shading(cell, 'E3F2FD')

    doc.add_page_break()

    # =========================================================================
    # CDC PATTERN
    # =========================================================================
    add_styled_heading(doc, '🔄 Change Data Capture (CDC) Pattern', 1)

    doc.add_paragraph(
        "Spread Eagle uses a staging-based CDC pattern for incremental data loads:"
    )

    cdc_steps = [
        ('1️⃣', 'EXTRACT: Pull new/changed records from API using date filters', 'source'),
        ('2️⃣', 'STAGE: Load into stg_* tables (TRUNCATE + INSERT)', 'storage'),
        ('3️⃣', 'UPSERT: INSERT ON CONFLICT to merge into main tables', 'transform'),
        ('4️⃣', 'CLEAN: Truncate staging tables after successful merge', 'python'),
        ('5️⃣', 'TRACK: load_date column records when each row was updated', 'output'),
    ]

    create_flowchart_table(doc, '🔄 CDC Workflow', cdc_steps)

    doc.add_paragraph()
    doc.add_paragraph(
        "This pattern ensures idempotent loads—running the same load twice produces "
        "the same result. The load_date column enables time-travel queries and debugging."
    )

    doc.add_page_break()

    # =========================================================================
    # ANALYTICS FEATURES
    # =========================================================================
    add_styled_heading(doc, '📈 Analytics & Features', 1)

    doc.add_paragraph(
        "The platform calculates key metrics for identifying betting opportunities:"
    )

    # ATS Analysis
    add_styled_heading(doc, '🎯 Against The Spread (ATS) Analysis', 2)

    ats_box = doc.add_paragraph()
    ats_box.add_run('Key Calculations:\n').bold = True
    ats_box.add_run('• actual_margin = home_score - away_score\n')
    ats_box.add_run('• ats_margin = actual_margin + spread\n')
    ats_box.add_run('• covered = (ats_margin > 0)\n')
    ats_box.add_run('• rolling_cover_rate = AVG(covered) over last N games\n')

    # Volatility Metrics
    add_styled_heading(doc, '📊 Volatility Metrics', 2)

    vol_box = doc.add_paragraph()
    vol_box.add_run('Consistency Indicators:\n').bold = True
    vol_box.add_run('• ats_volatility = STDDEV(ats_margin) — lower = more predictable\n')
    vol_box.add_run('• ou_volatility = STDDEV(total - over_under) — total consistency\n')
    vol_box.add_run('• consistency_score = 1 / (1 + STDDEV) — normalized 0-1 scale\n')

    # Rolling Windows
    add_styled_heading(doc, '📅 Rolling Window Features', 2)

    roll_box = doc.add_paragraph()
    roll_box.add_run('Time Windows (3, 5, 10 games):\n').bold = True
    roll_box.add_run('• rolling_avg_points_scored / allowed\n')
    roll_box.add_run('• rolling_avg_ats_margin\n')
    roll_box.add_run('• rolling_stddev_ats_margin\n')
    roll_box.add_run('• rolling_cover_rate\n')
    roll_box.add_run('• rolling_avg_pace / offensive_rating / defensive_rating\n')

    doc.add_page_break()

    # =========================================================================
    # FILE STRUCTURE
    # =========================================================================
    add_styled_heading(doc, '📁 Project Structure', 1)

    structure = """
spread_eagle/
├── 📁 infra/
│   └── terraform/app/          # AWS infrastructure as code
│       ├── providers.tf        # AWS provider config
│       ├── variables.tf        # Configurable settings
│       ├── vpc.tf              # Network setup
│       ├── s3.tf               # Data lake bucket
│       ├── rds.tf              # PostgreSQL database
│       ├── iam.tf              # Permissions
│       └── secrets.tf          # Credential storage
│
├── 📁 spread_eagle/
│   ├── config/
│   │   └── settings.py         # Pydantic settings
│   │
│   └── ingest/cbb/
│       ├── _common.py          # Shared utilities
│       ├── pull_games_full.py  # Full game load
│       ├── pull_lines_full.py  # Betting lines
│       ├── pull_team_stats_full.py
│       ├── pull_game_players_full.py
│       ├── pull_*_season_stats_full.py
│       ├── pull_conferences.py # Reference tables
│       ├── pull_venues.py
│       ├── pull_teams.py
│       ├── run_full_load.py    # Master orchestrator
│       └── generate_ddl.py     # Schema generator
│
├── 📁 data/cbb/
│   ├── raw/                    # Downloaded JSON/CSV
│   │   ├── games/
│   │   ├── lines/
│   │   ├── team_stats/
│   │   └── ...
│   └── ddl/
│       ├── create_tables.sql   # PostgreSQL DDL
│       └── upsert_from_staging.sql
│
├── 📁 dbt/                     # (future) Transform models
│   ├── models/
│   │   ├── staging/
│   │   ├── intermediate/
│   │   └── marts/
│   └── dbt_project.yml
│
└── .env                        # Environment variables
"""

    struct_para = doc.add_paragraph()
    struct_run = struct_para.add_run(structure)
    struct_run.font.name = 'Consolas'
    struct_run.font.size = Pt(9)

    doc.add_page_break()

    # =========================================================================
    # S3 STRUCTURE
    # =========================================================================
    add_styled_heading(doc, '📦 S3 Data Lake Structure', 1)

    s3_structure = """
s3://spread-eagle/
└── cbb/
    └── raw/
        ├── conferences/
        │   ├── conferences.json
        │   ├── conferences.csv
        │   └── conferences.parquet
        │
        ├── venues/
        │   └── venues.{json,csv,parquet}
        │
        ├── teams/
        │   └── teams.{json,csv,parquet}
        │
        ├── games/
        │   ├── games_2022.json
        │   ├── games_2023.json
        │   ├── games_2024.json
        │   ├── games_2025.json
        │   ├── games_2026.json
        │   ├── games_2022_2026.csv      # Consolidated
        │   └── games_2022_2026.parquet
        │
        ├── lines/
        │   ├── lines_{year}.json
        │   ├── lines_2022_2026.csv
        │   └── lines_2022_2026.parquet
        │
        ├── team_stats/
        │   └── team_stats_*.{json,csv,parquet}
        │
        ├── game_players/
        │   └── game_players_*.{json,csv,parquet}
        │
        ├── team_season_stats/
        │   └── team_season_stats_*.{json,csv,parquet}
        │
        └── player_season_stats/
            └── player_season_stats_*.{json,csv,parquet}
"""

    s3_para = doc.add_paragraph()
    s3_run = s3_para.add_run(s3_structure)
    s3_run.font.name = 'Consolas'
    s3_run.font.size = Pt(9)

    doc.add_page_break()

    # =========================================================================
    # COMMANDS REFERENCE
    # =========================================================================
    add_styled_heading(doc, '⌨️ Commands Reference', 1)

    add_styled_heading(doc, 'Full Data Load', 2)
    cmd1 = doc.add_paragraph()
    cmd1.add_run('# Run all ingestion scripts\n').font.name = 'Consolas'
    cmd1.add_run('python -m spread_eagle.ingest.cbb.run_full_load').font.name = 'Consolas'

    add_styled_heading(doc, 'Individual Endpoints', 2)
    cmd2 = doc.add_paragraph()
    cmds = [
        'python -m spread_eagle.ingest.cbb.pull_conferences',
        'python -m spread_eagle.ingest.cbb.pull_venues',
        'python -m spread_eagle.ingest.cbb.pull_teams',
        'python -m spread_eagle.ingest.cbb.pull_games_full',
        'python -m spread_eagle.ingest.cbb.pull_lines_full',
        'python -m spread_eagle.ingest.cbb.pull_team_stats_full',
        'python -m spread_eagle.ingest.cbb.pull_game_players_full',
    ]
    for cmd in cmds:
        cmd2.add_run(cmd + '\n').font.name = 'Consolas'

    add_styled_heading(doc, 'Database Operations', 2)
    cmd3 = doc.add_paragraph()
    cmd3.add_run('# Generate DDL from CSV schemas\n').font.name = 'Consolas'
    cmd3.add_run('python -m spread_eagle.ingest.cbb.generate_ddl\n\n').font.name = 'Consolas'
    cmd3.add_run('# Run DDL on RDS\n').font.name = 'Consolas'
    cmd3.add_run('python -m spread_eagle.ingest.cbb.run_ddl\n\n').font.name = 'Consolas'
    cmd3.add_run('# Or use psql directly\n').font.name = 'Consolas'
    cmd3.add_run('psql -h <host> -U postgres -d postgres -f data/cbb/ddl/create_tables.sql').font.name = 'Consolas'

    add_styled_heading(doc, 'Infrastructure', 2)
    cmd4 = doc.add_paragraph()
    cmd4.add_run('# Deploy AWS infrastructure\n').font.name = 'Consolas'
    cmd4.add_run('cd infra/terraform/app\n').font.name = 'Consolas'
    cmd4.add_run('terraform init\n').font.name = 'Consolas'
    cmd4.add_run('terraform plan\n').font.name = 'Consolas'
    cmd4.add_run('terraform apply\n\n').font.name = 'Consolas'
    cmd4.add_run('# Stop RDS to save money\n').font.name = 'Consolas'
    cmd4.add_run('aws rds stop-db-instance --db-instance-identifier spread-eagle-db --region us-east-2\n\n').font.name = 'Consolas'
    cmd4.add_run('# Start RDS\n').font.name = 'Consolas'
    cmd4.add_run('aws rds start-db-instance --db-instance-identifier spread-eagle-db --region us-east-2').font.name = 'Consolas'

    doc.add_page_break()

    # =========================================================================
    # FUTURE ROADMAP
    # =========================================================================
    add_styled_heading(doc, '🚀 Future Roadmap', 1)

    roadmap = [
        ('Phase 1: Foundation', [
            '✅ AWS Infrastructure (Terraform)',
            '✅ Data Ingestion Scripts (Python)',
            '✅ S3 Data Lake',
            '✅ PostgreSQL Schema (DDL)',
            '⏳ dbt Staging Models',
        ]),
        ('Phase 2: Analytics', [
            '⏳ Rolling Window Calculations',
            '⏳ ATS Performance Metrics',
            '⏳ Volatility Scoring',
            '⏳ Teaser Value Analysis',
        ]),
        ('Phase 3: Visualization', [
            '⏳ Dashboard (Streamlit/Dash)',
            '⏳ Team Comparison Tools',
            '⏳ Daily Line Movement Alerts',
        ]),
        ('Phase 4: Expansion', [
            '⏳ College Football (CFB)',
            '⏳ NFL/NBA Integration',
            '⏳ Live Odds API',
            '⏳ ML Prediction Models',
        ]),
    ]

    for phase, items in roadmap:
        add_styled_heading(doc, phase, 2)
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(item)

    # =========================================================================
    # SAVE
    # =========================================================================
    output_dir = Path('docs')
    output_dir.mkdir(exist_ok=True)

    output_path = output_dir / 'Spread_Eagle_Architecture.docx'
    doc.save(output_path)
    print(f"Created: {output_path}")
    print(f"File size: {output_path.stat().st_size:,} bytes")


if __name__ == '__main__':
    main()

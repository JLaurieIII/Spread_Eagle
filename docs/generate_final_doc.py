"""
Generate final Spread Eagle documentation with embedded images.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def main():
    doc = Document()

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    for _ in range(3):
        doc.add_paragraph()

    title = doc.add_heading('', 0)
    title_run = title.add_run('SPREAD EAGLE')
    title_run.font.size = Pt(48)
    title_run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('Sports Betting Analytics Platform')
    subtitle_run.font.size = Pt(24)
    subtitle_run.font.color.rgb = RGBColor(0x5c, 0x6b, 0x7a)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    tagline = doc.add_paragraph()
    tagline_run = tagline.add_run('Find the Tail. Ride the Edge.')
    tagline_run.font.size = Pt(16)
    tagline_run.font.italic = True
    tagline_run.font.color.rgb = RGBColor(0x6c, 0x75, 0x7d)
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(3):
        doc.add_paragraph()

    # Add main architecture image
    arch_img = Path('docs/Spread_Eagle_Architecture.png')
    if arch_img.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(arch_img), width=Inches(5.5))

    doc.add_page_break()

    # =========================================================================
    # EXECUTIVE SUMMARY
    # =========================================================================
    heading = doc.add_heading('Executive Summary', 1)
    heading.runs[0].font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    doc.add_paragraph(
        "Spread Eagle is a modern data analytics platform designed to identify "
        "low-volatility betting opportunities in college basketball. By analyzing "
        "historical betting lines, game statistics, and team performance metrics, "
        "the platform identifies teams that consistently play close to betting lines - "
        "creating valuable 'tail' opportunities for teaser bets."
    )

    doc.add_paragraph()

    # Mission statement box
    mission = doc.add_paragraph()
    mission.add_run('Mission: ').bold = True
    mission.add_run(
        "Find teams with predictable performance against the spread, "
        "enabling data-driven betting decisions with reduced variance."
    )

    doc.add_paragraph()

    # Key numbers
    heading2 = doc.add_heading('By The Numbers', 2)
    heading2.runs[0].font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    metrics_table = doc.add_table(rows=2, cols=5)

    metrics = [
        ('31,056', 'Games'),
        ('49,153', 'Betting Lines'),
        ('55,062', 'Team Stats'),
        ('222,594', 'Player Stats'),
        ('5', 'Seasons'),
    ]

    for i, (value, label) in enumerate(metrics):
        cell = metrics_table.cell(0, i)
        cell.text = value
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '1a365d')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        cell2 = metrics_table.cell(1, i)
        cell2.text = label
        cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell2, 'E3F2FD')

    doc.add_page_break()

    # =========================================================================
    # ARCHITECTURE DIAGRAM
    # =========================================================================
    heading = doc.add_heading('System Architecture', 1)
    heading.runs[0].font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    doc.add_paragraph(
        "Spread Eagle follows a modern ELT (Extract, Load, Transform) architecture "
        "with clear separation between data ingestion, storage, and transformation layers."
    )

    # Add architecture image
    arch_img = Path('docs/Spread_Eagle_Architecture.png')
    if arch_img.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(arch_img), width=Inches(6.5))

    doc.add_page_break()

    # =========================================================================
    # DATA FLOW
    # =========================================================================
    heading = doc.add_heading('Data Pipeline Flow', 1)
    heading.runs[0].font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    doc.add_paragraph(
        "Data flows through six distinct stages, each optimized for its purpose:"
    )

    # Add flow diagram
    flow_img = Path('docs/Spread_Eagle_DataFlow.png')
    if flow_img.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(flow_img), width=Inches(6.5))

    doc.add_paragraph()

    # Flow stages
    stages = [
        ('1. API EXTRACTION', 'College Basketball Data API provides REST endpoints with JSON responses. The API has a 3,000-record limit per request.'),
        ('2. PYTHON INGESTION', 'Custom scripts use date-range pagination to overcome API limits. Data is fetched monthly, deduplicated, and uploaded to S3.'),
        ('3. S3 DATA LAKE', 'Raw data stored as JSON (by season), CSV (consolidated), and Parquet (optimized). Lifecycle policies archive old data.'),
        ('4. POSTGRESQL LOADING', 'Staging tables receive new data via COPY command. CDC pattern (INSERT ON CONFLICT) merges into main tables.'),
        ('5. DBT TRANSFORMATION', 'SQL models calculate rolling statistics, ATS performance, volatility metrics. Incremental processing for efficiency.'),
        ('6. ANALYTICS OUTPUT', 'Mart tables ready for dashboards, reports, and automated alerts.'),
    ]

    for title, desc in stages:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        p.add_run(f'\n{desc}')

    doc.add_page_break()

    # =========================================================================
    # TECHNOLOGY STACK
    # =========================================================================
    heading = doc.add_heading('Technology Stack', 1)
    heading.runs[0].font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    # Add tech stack image
    tech_img = Path('docs/Spread_Eagle_TechStack.png')
    if tech_img.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(tech_img), width=Inches(5.5))

    doc.add_paragraph()

    # Tool details
    tools = [
        ('Python', 'Core ingestion and orchestration', ['requests - API calls', 'boto3 - AWS SDK', 'pandas - Data manipulation', 'psycopg2 - PostgreSQL driver']),
        ('AWS S3', 'Scalable data lake storage', ['JSON raw files', 'CSV for analysis', 'Parquet for performance', 'Lifecycle policies']),
        ('PostgreSQL (RDS)', 'Relational data warehouse', ['cbb schema', 'Staging tables (stg_*)', 'Fact and dimension tables', 'CDC with load_date tracking']),
        ('dbt', 'Data transformation framework', ['Modular SQL models', 'Jinja templating', 'Incremental processing', 'Built-in testing']),
        ('Terraform', 'Infrastructure as code', ['VPC and networking', 'RDS provisioning', 'S3 bucket creation', 'IAM policies']),
    ]

    for name, desc, items in tools:
        heading2 = doc.add_heading(name, 2)
        heading2.runs[0].font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
        doc.add_paragraph(desc)
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(item)

    doc.add_page_break()

    # =========================================================================
    # DATABASE SCHEMA
    # =========================================================================
    heading = doc.add_heading('Database Schema', 1)
    heading.runs[0].font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    doc.add_paragraph(
        "The database is organized into reference tables (dimensions) and transactional tables (facts):"
    )

    # Schema table
    tables_data = [
        ('Table', 'Type', 'Primary Key', 'Records'),
        ('conferences', 'Reference', 'id', '34'),
        ('venues', 'Reference', 'id', '979'),
        ('teams', 'Reference', 'id', '1,515'),
        ('games', 'Fact', 'id', '31,056'),
        ('betting_lines', 'Fact', 'game_id, provider', '49,153'),
        ('game_team_stats', 'Fact', 'game_id, team_id', '55,062'),
        ('game_player_stats', 'Fact', 'game_id, athlete_id', '222,594'),
        ('team_season_stats', 'Fact', 'team_id, season', '3,515'),
        ('player_season_stats', 'Fact', 'athlete_id, team_id, season', '48,551'),
    ]

    table = doc.add_table(rows=len(tables_data), cols=4)
    table.style = 'Table Grid'

    for i, row_data in enumerate(tables_data):
        for j, cell_data in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = cell_data
            if i == 0:
                set_cell_shading(cell, '1a365d')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # =========================================================================
    # ANALYTICS FEATURES
    # =========================================================================
    heading = doc.add_heading('Analytics Features', 1)
    heading.runs[0].font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    doc.add_paragraph(
        "The platform calculates key metrics for identifying betting opportunities:"
    )

    # ATS
    heading2 = doc.add_heading('Against The Spread (ATS) Analysis', 2)
    heading2.runs[0].font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    formulas = [
        'actual_margin = home_score - away_score',
        'ats_margin = actual_margin + spread',
        'covered = (ats_margin > 0)',
        'push = (ats_margin = 0)',
    ]
    for f in formulas:
        p = doc.add_paragraph()
        run = p.add_run(f)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)

    # Volatility
    heading2 = doc.add_heading('Volatility Metrics', 2)
    heading2.runs[0].font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    doc.add_paragraph(
        "Lower volatility = more predictable performance against the spread:"
    )

    vol_formulas = [
        'ats_volatility = STDDEV(ats_margin) over last N games',
        'consistency_score = 1 / (1 + STDDEV(ats_margin))',
        'ou_volatility = STDDEV(total_points - over_under)',
    ]
    for f in vol_formulas:
        p = doc.add_paragraph()
        run = p.add_run(f)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)

    # Rolling windows
    heading2 = doc.add_heading('Rolling Window Features', 2)
    heading2.runs[0].font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    doc.add_paragraph('Calculated over 3, 5, and 10 game windows:')

    rolling = [
        'rolling_avg_points_scored / allowed',
        'rolling_avg_ats_margin',
        'rolling_stddev_ats_margin',
        'rolling_cover_rate',
        'rolling_avg_pace',
        'rolling_avg_offensive_rating / defensive_rating',
    ]
    for r in rolling:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(r)

    doc.add_page_break()

    # =========================================================================
    # COMMANDS REFERENCE
    # =========================================================================
    heading = doc.add_heading('Commands Reference', 1)
    heading.runs[0].font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    # Full load
    heading2 = doc.add_heading('Full Data Load', 2)
    heading2.runs[0].font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    p = doc.add_paragraph()
    p.add_run('python -m spread_eagle.ingest.cbb.run_full_load').font.name = 'Consolas'

    # Individual
    heading2 = doc.add_heading('Individual Endpoints', 2)
    heading2.runs[0].font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    cmds = [
        'python -m spread_eagle.ingest.cbb.pull_conferences',
        'python -m spread_eagle.ingest.cbb.pull_venues',
        'python -m spread_eagle.ingest.cbb.pull_teams',
        'python -m spread_eagle.ingest.cbb.pull_games_full',
        'python -m spread_eagle.ingest.cbb.pull_lines_full',
        'python -m spread_eagle.ingest.cbb.pull_team_stats_full',
        'python -m spread_eagle.ingest.cbb.pull_game_players_full',
        'python -m spread_eagle.ingest.cbb.pull_team_season_stats_full',
        'python -m spread_eagle.ingest.cbb.pull_player_season_stats_full',
    ]
    for cmd in cmds:
        p = doc.add_paragraph()
        p.add_run(cmd).font.name = 'Consolas'

    # Database
    heading2 = doc.add_heading('Database Operations', 2)
    heading2.runs[0].font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    db_cmds = [
        '# Generate DDL from CSV schemas',
        'python -m spread_eagle.ingest.cbb.generate_ddl',
        '',
        '# Run DDL on RDS',
        'python -m spread_eagle.ingest.cbb.run_ddl',
    ]
    for cmd in db_cmds:
        p = doc.add_paragraph()
        p.add_run(cmd).font.name = 'Consolas'

    # Infrastructure
    heading2 = doc.add_heading('Infrastructure Management', 2)
    heading2.runs[0].font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    infra_cmds = [
        'cd infra/terraform/app',
        'terraform init',
        'terraform plan',
        'terraform apply',
        '',
        '# Stop RDS to save money',
        'aws rds stop-db-instance --db-instance-identifier spread-eagle-db',
        '',
        '# Start RDS',
        'aws rds start-db-instance --db-instance-identifier spread-eagle-db',
    ]
    for cmd in infra_cmds:
        p = doc.add_paragraph()
        p.add_run(cmd).font.name = 'Consolas'

    doc.add_page_break()

    # =========================================================================
    # ROADMAP
    # =========================================================================
    heading = doc.add_heading('Roadmap', 1)
    heading.runs[0].font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    roadmap = [
        ('Phase 1: Foundation (Complete)', [
            'AWS Infrastructure (Terraform)',
            'Data Ingestion Scripts (Python)',
            'S3 Data Lake',
            'PostgreSQL Schema (DDL)',
        ]),
        ('Phase 2: Analytics (In Progress)', [
            'dbt Staging Models',
            'Rolling Window Calculations',
            'ATS Performance Metrics',
            'Volatility Scoring',
        ]),
        ('Phase 3: Visualization', [
            'Dashboard (Streamlit/Dash)',
            'Team Comparison Tools',
            'Daily Line Movement Alerts',
        ]),
        ('Phase 4: Expansion', [
            'College Football (CFB)',
            'NFL/NBA Integration',
            'Live Odds API',
            'ML Prediction Models',
        ]),
    ]

    for phase, items in roadmap:
        heading2 = doc.add_heading(phase, 2)
        heading2.runs[0].font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(item)

    # =========================================================================
    # SAVE
    # =========================================================================
    output_path = Path('docs/Spread_Eagle_Documentation.docx')
    doc.save(output_path)
    print(f"Created: {output_path}")
    print(f"Size: {output_path.stat().st_size:,} bytes")


if __name__ == '__main__':
    main()
